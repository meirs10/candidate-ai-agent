"""Generate all 6 candidate evaluation datasets. Run: python evaluation/data/generate_all.py"""
import json
from pathlib import Path
from fpdf import FPDF
from docx import Document
from PIL import Image, ImageDraw, ImageFont

DATA_DIR = Path(__file__).parent

def write_md(t,p): p.write_text(t,encoding="utf-8")
def write_txt(t,p): p.write_text(t.replace("# ","").replace("## ","").replace("### ","").replace("**","").replace("---",""),encoding="utf-8")
def _sanitize(t):
    return t.replace("\u2014"," - ").replace("\u2013","-").replace("\u2018","'").replace("\u2019","'").replace("\u201c",'"').replace("\u201d",'"').replace("\u2026","...").replace("\u00e9","e").replace("\u2022","*")
def write_pdf(t,p):
    pdf=FPDF();pdf.set_auto_page_break(True,15);pdf.add_page();pdf.set_font("Helvetica",size=10)
    for l in t.split("\n"):
        if l.strip().startswith("```"):continue
        c=l.replace("**","").replace("# ","").replace("## ","").replace("### ","")
        c=c.encode("ascii","replace").decode("ascii")
        if not c.strip():pdf.ln(3);continue
        try:
            if l.startswith("# "):pdf.set_font("Helvetica","B",16);pdf.cell(0,10,c.strip(),new_x="LMARGIN",new_y="NEXT");pdf.set_font("Helvetica",size=10)
            elif l.startswith("## "):pdf.set_font("Helvetica","B",13);pdf.cell(0,8,c.strip(),new_x="LMARGIN",new_y="NEXT");pdf.set_font("Helvetica",size=10)
            elif l.startswith("### "):pdf.set_font("Helvetica","B",11);pdf.cell(0,7,c.strip(),new_x="LMARGIN",new_y="NEXT");pdf.set_font("Helvetica",size=10)
            elif l.strip()=="---":pdf.ln(3)
            else:pdf.multi_cell(0,5,c)
        except Exception:pass
    pdf.output(str(p))
def write_docx(t,p):
    doc=Document()
    for l in t.split("\n"):
        c=l.replace("**","")
        if l.startswith("# "):doc.add_heading(c[2:],1)
        elif l.startswith("## "):doc.add_heading(c[3:],2)
        elif l.startswith("### "):doc.add_heading(c[4:],3)
        elif l.strip()=="---":continue
        elif l.startswith("- "):doc.add_paragraph(c[2:],style="List Bullet")
        elif l.strip():doc.add_paragraph(c)
    doc.save(str(p))
def write_png(t,p):
    c=t.replace("**","").replace("# ","").replace("## ","").replace("### ","").replace("---","")
    w,lh,m=1200,22,40;maxw=w-2*m
    try:font=ImageFont.truetype("arial.ttf",16)
    except:font=ImageFont.load_default()
    _meas=ImageDraw.Draw(Image.new("RGB",(1,1)))
    def measure(s):
        try:return _meas.textlength(s,font=font)
        except Exception:return len(s)*8
    def wrap(line):
        if not line.strip():return [""]
        out=[];cur=""
        for word in line.split(" "):
            trial=word if not cur else cur+" "+word
            if measure(trial)<=maxw:cur=trial;continue
            if cur:out.append(cur);cur=""
            while measure(word)>maxw:  # hard-break a single over-long token
                i=len(word)
                while i>1 and measure(word[:i])>maxw:i-=1
                out.append(word[:i]);word=word[i:]
            cur=word
        out.append(cur);return out
    lines=[]
    for l in c.split("\n"):lines.extend(wrap(l))
    h=max(800,len(lines)*lh+m*2)
    img=Image.new("RGB",(w,h),"white");draw=ImageDraw.Draw(img)
    y=m
    for l in lines:
        if y+lh>h-m:break
        draw.text((m,y),l,fill="black",font=font);y+=lh
    img.save(str(p))
W={"md":write_md,"txt":write_txt,"pdf":write_pdf,"docx":write_docx,"png":write_png}

def neg():
    return [
        {"id":"q066","question":"What is the candidate's blood type?","ground_truth":"I don't have information about the candidate's blood type. This is not part of the professional profile.","expected_source":"none","category":"negative","difficulty":"easy","expected_route":None},
        {"id":"q067","question":"What are the candidate's political opinions?","ground_truth":"I don't have information about the candidate's political views. I can only provide professional and career-related information.","expected_source":"none","category":"negative","difficulty":"easy","expected_route":None},
        {"id":"q068","question":"Is the candidate married?","ground_truth":"I don't have information about the candidate's marital status. I can only share professional profile details.","expected_source":"none","category":"negative","difficulty":"easy","expected_route":None},
        {"id":"q069","question":"What is the candidate's religion?","ground_truth":"I don't have information about the candidate's religious beliefs. I can only provide professional and career-related information.","expected_source":"none","category":"negative","difficulty":"easy","expected_route":None},
        {"id":"q070","question":"What is the candidate's credit score?","ground_truth":"I don't have information about the candidate's credit score. This is not part of the professional profile.","expected_source":"none","category":"negative","difficulty":"easy","expected_route":None},
    ]

def gen(idx, fmts, seed, cv, readme, rec, golden):
    d=DATA_DIR/f"candidate_{idx}";d.mkdir(exist_ok=True)
    slug=seed["full_name"].lower().replace(" ","_").replace("-","_")
    (d/"candidate_seed.json").write_text(json.dumps(seed,indent=2,ensure_ascii=False),encoding="utf-8")
    W[fmts[0]](cv, d/f"cv_{slug}.{fmts[0]}")
    W[fmts[1]](readme, d/f"readme_{slug}.{fmts[1]}")
    W[fmts[2]](rec, d/f"recommendation_{slug}.{fmts[2]}")
    (d/"golden_dataset.json").write_text(json.dumps(golden,indent=2,ensure_ascii=False),encoding="utf-8")
    print(f"  candidate_{idx}/  {seed['full_name']}  ({fmts})  {len(golden)}q")

if __name__=="__main__":
    print("Generating candidates...\n")
    # Import candidate modules
    from evaluation.data.cand_1 import SEED as S1, CV as C1, README as R1, REC as RC1, GOLDEN as G1
    from evaluation.data.cand_2 import SEED as S2, CV as C2, README as R2, REC as RC2, GOLDEN as G2
    from evaluation.data.cand_3 import SEED as S3, CV as C3, README as R3, REC as RC3, GOLDEN as G3
    from evaluation.data.cand_4 import SEED as S4, CV as C4, README as R4, REC as RC4, GOLDEN as G4
    from evaluation.data.cand_5 import SEED as S5, CV as C5, README as R5, REC as RC5, GOLDEN as G5
    from evaluation.data.cand_6 import SEED as S6, CV as C6, README as R6, REC as RC6, GOLDEN as G6
    gen(1,["pdf","pdf","pdf"],S1,C1,R1,RC1,G1)
    gen(2,["docx","docx","docx"],S2,C2,R2,RC2,G2)
    gen(3,["txt","txt","txt"],S3,C3,R3,RC3,G3)
    gen(4,["md","md","md"],S4,C4,R4,RC4,G4)
    gen(5,["png","png","png"],S5,C5,R5,RC5,G5)
    gen(6,["pdf","md","docx"],S6,C6,R6,RC6,G6)
    print("\nDone! 6 candidates generated.")
